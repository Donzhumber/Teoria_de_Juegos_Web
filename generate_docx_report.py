import os
from docx import Document
from docx.shared import Inches

def create_report_docx():
    doc = Document()
    doc.add_heading('Informe de Documentos Confecámaras', 0)

    # Content from the markdown report
    # (Manually adding it for simplicity and formatting control)
    doc.add_heading('Análisis y Consecuencias', level=1)
    doc.add_paragraph(
        "Este informe resume el contenido de los documentos encontrados en Gmail y analiza sus "
        "implicaciones legales y administrativas bajo el régimen de Garantías Mobiliarias en Colombia (Ley 1676 de 2013)."
    )

    doc.add_heading('Documentos Identificados', level=2)
    
    # Document 1
    doc.add_heading('1. Registro de Terminación de la Ejecución', level=3)
    doc.add_paragraph("Fecha: 11/02/2026 10:04:51")
    doc.add_paragraph("Folio Electrónico: 20200730000021700")
    doc.add_paragraph("Causal: Vencimiento del plazo para la ejecución de la garantía.")
    
    img1_path = '/Users/zhumber/.gemini/antigravity/brain/6b4db796-66fc-4ff8-aed3-1d3e42c5cb48/pdf_2_preview_1773662442658.png'
    if os.path.exists(img1_path):
        doc.add_picture(img1_path, width=Inches(6))
        doc.add_paragraph("Vista previa del Registro de Terminación")

    # Document 2
    doc.add_heading('2. Formulario de Registro de Ejecución', level=3)
    doc.add_paragraph("Fecha: 17/02/2026 14:35:38")
    doc.add_paragraph("Folio Electrónico: 20200730000021700")
    doc.add_paragraph("Estado: Proceso de ejecución reiniciado por Banco Davivienda.")
    
    img2_path = '/Users/zhumber/.gemini/antigravity/brain/6b4db796-66fc-4ff8-aed3-1d3e42c5cb48/pdf_1_preview_1773662338630.png'
    if os.path.exists(img2_path):
        doc.add_picture(img2_path, width=Inches(6))
        doc.add_paragraph("Vista previa del Formulario de Registro de Ejecución")

    doc.add_heading('Análisis y Recomendaciones', level=2)
    doc.add_paragraph(
        "El registro del 17 de febrero reactiva formalmente el proceso de cobro. Se recomienda "
        "revisar el estado de la obligación con el Banco Davivienda para evitar la aprehensión de bienes."
    )

    output_path = '/Users/zhumber/Library/CloudStorage/OneDrive-Personal/Documents/Antigravity/Varios/Confecamaras_Docs/Informe_Confecamaras.docx'
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Reporte guardado en: {output_path}")

if __name__ == "__main__":
    create_report_docx()
